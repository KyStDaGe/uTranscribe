from docx import Document
from fpdf import FPDF

import threading
import pystray
from pystray import MenuItem as item
import customtkinter
import subprocess
import tkinter
from tkinter import filedialog, Label
import os
from PIL import Image
import speech_recognition as sr
import os
from recognize_text_from_image import interface

r = sr.Recognizer()

class TranscribeThread(threading.Thread):
    def __init__(self, event, parent):
        threading.Thread.__init__(self)
        self.event = event
        self.parent = parent

    def run(self):
        App.withdraw(self.parent)
        directory = os.path.join(os.path.dirname(os.path.realpath(__file__)))
        checkfile = os.path.isfile(os.path.join(directory, "out.wav"))
        splash = Splash(self.parent)
        if checkfile:
            a = Transcribe.transcribe_audio()
            os.remove(os.path.join(directory, "out.wav"))
        splash.destroy()
        App.deiconify(self.parent)
        App.savetextbox(self.parent, a)

class OCRThread(threading.Thread):
    def __init__(self, event, parent):
        threading.Thread.__init__(self)
        self.event = event
        self.parent = parent
        
    def run(self):
        App.withdraw(self.parent)
        directory = os.path.join(os.path.dirname(os.path.realpath(__file__)))
        checkfile = os.path.isfile(os.path.join(directory, "capture.jpg"))
        splash = Splash(self.parent)
        
        if checkfile:
            predictor = interface.Predictor()
            prediction = predictor.predict(os.path.join(directory, "capture.jpg"))
            
        splash.destroy()
        App.deiconify(self.parent)
        App.savetextbox2(self.parent,prediction)

class Splash(tkinter.Toplevel):
    def __init__(self, parent):
        tkinter.Toplevel.__init__(self, parent)
        def hide_window():
            pass
        self.title("Splash")
        splash_label = Label(self, text="Loading..", font=18)
        splash_label.pack(padx=20, pady=10)
        self.geometry("200x100")
        self.minsize(width=200,height=100)
        self.maxsize(width=200,height=100)
        self.protocol("WM_DELETE_WINDOW",hide_window)
        progressbar = customtkinter.CTkProgressBar(master=self)
        progressbar.pack(padx=20, pady=10)
        progressbar.configure(mode="indeterminate")
        progressbar.start()
        
class Transcribe():
    def transcribe_audio():
        directory = os.path.join(os.path.dirname(os.path.realpath(__file__)))
        dir = os.path.join(directory,"parts")
        temp = 0
        count = 0
        loading = 0
        for path in os.listdir(dir):
            if os.path.isfile(os.path.join(dir, path)):
                count += 1
        print(count)    
        object = os.scandir(dir)
        whole_text = ""
        for n in object:
            if n.is_file():
                chunk_filename = os.path.join(dir,n.name)
            with sr.AudioFile(chunk_filename) as source:
                audio_listened = r.record(source)
            # Convert to text
                try:
                    text = r.recognize_google(audio_listened)
                except sr.RequestError as e:
                    text = r.recognize_sphinx(audio_listened)
                    text = f"{text.capitalize()} "
                    print(chunk_filename, ":Done.")
                    whole_text += text
                    
                except sr.UnknownValueError as e:
                    print("Error:", str(e))
                else:
                    text = f"{text.capitalize()} "
                    print(chunk_filename, ":Done.")
                    whole_text += text
                    loading += 1
                    temp = (loading/count)*100//1
                    print(temp)
                    
            os.remove(chunk_filename)
        return whole_text
    
class App(customtkinter.CTk):
    def __init__(self, *args, **kwargs):
        def snip():
            directory = os.path.join(os.path.dirname(os.path.realpath(__file__)))
            snippath = os.path.join(directory,"snipping_tool.exe")
            subprocess.run(snippath)
            event = threading.Event()
            transcribe_thread = OCRThread(event, self)
            transcribe_thread.start()
            
            
        def convert():
            directory = os.path.join(os.path.dirname(os.path.realpath(__file__)))
            ffmpegpath =os.path.join(directory,"ffmpeg/bin/ffmpeg.exe")
            filename = filedialog.askopenfilename(title="Select a File", filetypes=([("Audio or Video Files", "*.wav , *.mp4 , *.mp3")]))
            if filename is None:
                return
            else:
                patha = filename
                filepath = "Selected file: " + filename
                self.filelabel = customtkinter.CTkLabel(self.speech_frame, text=filepath,
                font=customtkinter.CTkFont(size=15, weight="normal"))
                self.filelabel.grid(row=0, column=0, padx=(20, 20), pady=(20, 20), sticky="nsew")
            
                checkfile = os.path.isfile(os.path.join(directory,"out.wav"))
            
                if checkfile == True:
                    os.remove(os.path.join(directory,"out.wav"))
                    subprocess.call([ffmpegpath,'-i',patha,'-acodec','pcm_s16le','-ac','1','-ar','16000','out.wav'])
                    split()    
                else:
                    subprocess.call([ffmpegpath,'-i',patha,'-acodec','pcm_s16le','-ac','1','-ar','16000','out.wav'])  
                    split()  
                                
        
        def split():
            directory = os.path.join(os.path.dirname(os.path.realpath(__file__)))
            ffmpegpath =os.path.join(directory,"ffmpeg/bin/ffmpeg.exe") 
            checkfile = os.path.isfile(os.path.join(directory,"out.wav"))
            if checkfile == True:
                subprocess.call([ffmpegpath,'-i','out.wav','-f','segment','-segment_time','30','-c','copy','parts/output%3d.wav'])
            
        def transcribe():
            directory = os.path.join(os.path.dirname(os.path.realpath(__file__)))
            checkfile = os.path.isfile(os.path.join(directory, "out.wav"))
            if checkfile:
                event = threading.Event()
                transcribe_thread = TranscribeThread(event, self)
                transcribe_thread.start()
                

            
        def save1txt():
            files = [('Text Document', '*.txt')]
            f = filedialog.asksaveasfile(mode='w', defaultextension=files, filetypes=files)
            if f is None:
                return
            
            text2save = self.speech_frame_textbox.get(0.0,"end")
            f.write(text2save)
            f.close()
        
        def save1docx():
            document = Document()
            
            files = [('Word Document', '*.docx')]
            f = filedialog.asksaveasfile(mode='w', defaultextension=files, filetypes=files)
            if f is None:
                return
            name = f.name
            basename = os.path.basename(name)
            path = os.path.dirname(name)
            
            text2save = self.speech_frame_textbox.get(0.0,"end")
            document.add_paragraph(text2save)
            document.save(path + "/" + basename)
            f.close()    
        
        def save1pdf():
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('arial','B',13.0)
            files = [('PDF File', '*.pdf')]
            f = filedialog.asksaveasfile(mode='w', defaultextension=files, filetypes=files)
            if f is None:
                return
            name = f.name
            basename = os.path.basename(name)
            path = os.path.dirname(name)
            text2save = self.speech_frame_textbox.get(0.0,"end")
            pdf.multi_cell(0, 5, text2save, 0,'J', False)
            pdf.output(path + "/" + basename)
            f.close()    
        
        def save2():
            files = [('Text Document', '*.txt')]
            f = filedialog.asksaveasfile(mode='w', defaultextension=files, filetypes=files)
            text2save = self.ocr_frame_textbox.get(0.0,"end")
            if f is None:
                return
            
            f.write(text2save)
            f.close()
        
        def save2docx():
            document = Document()
            
            files = [('Word Document', '*.docx')]
            f = filedialog.asksaveasfile(mode='w', defaultextension=files, filetypes=files)
            if f is None:
                return
            name = f.name
            basename = os.path.basename(name)
            path = os.path.dirname(name)
            
            text2save = self.ocr_frame_textbox.get(0.0,"end")
            document.add_paragraph(text2save)
            document.save(path + "/" + basename)
            f.close()        
        
        def save2pdf():
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('arial','B',13.0)
            files = [('PDF File', '*.pdf')]
            f = filedialog.asksaveasfile(mode='w', defaultextension=files, filetypes=files)
            if f is None:
                return
            name = f.name
            basename = os.path.basename(name)
            path = os.path.dirname(name)
            text2save = self.ocr_frame_textbox.get(0.0,"end")
            pdf.multi_cell(0, 5, text2save, 0,'J', False)
            pdf.output(path + "/" + basename)
            f.close()
        def hide_window():
            self.withdraw()
            image=Image.open("favicon.ico")
            menu=(item('Close uTranscribe', quit_window), item('Open uTranscribe', show_window))
            icon=pystray.Icon("name", image, "uTranscribe", menu)
            icon.run()
        def quit_window(icon, item):
            icon.stop()
            self.destroy()
            
        def show_window(icon, item):
            icon.stop()
            self.after(1,self.deiconify)
                    
        super().__init__(*args, **kwargs)
        self.title("uTranscribe")
        self.geometry("1000x680+0+0")
        self.minsize(width=900,height=620)
        self.maxsize(width=1280,height=720)
        self.iconbitmap("favicon.ico")
        self.protocol('WM_DELETE_WINDOW', hide_window)
        # set grid layout 1x2
        self.grid_rowconfigure((0,1,2), weight=1)
        self.grid_columnconfigure(1, weight=1)
        
        # load images with light and dark mode image
        image_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), "test_images")
        self.logo_image = customtkinter.CTkImage(Image.open(os.path.join(image_path, "logo_single.png")), size=(55, 50))
        self.large_test_image = customtkinter.CTkImage(Image.open(os.path.join(image_path, "large_test_image.png")), size=(500, 150))
        self.image_icon_image = customtkinter.CTkImage(Image.open(os.path.join(image_path, "image_icon_light.png")), size=(20, 20))
        
        self.home_image = customtkinter.CTkImage(light_image=Image.open(os.path.join(image_path, "home_dark.png")),
                                                     dark_image=Image.open(os.path.join(image_path, "home_light.png")), size=(20, 20))
        self.speech_image = customtkinter.CTkImage(light_image=Image.open(os.path.join(image_path, "speech_light.png")),
                                                 dark_image=Image.open(os.path.join(image_path, "speech_dark.png")), size=(20, 20))
        self.ocr_image = customtkinter.CTkImage(light_image=Image.open(os.path.join(image_path, "ocr_dark.png")),
                                                 dark_image=Image.open(os.path.join(image_path, "ocr_light.png")), size=(15, 20))
       
        # create navigation frame
        self.navigation_frame = customtkinter.CTkFrame(self, corner_radius=0)
        self.navigation_frame.grid(row=0, column=0, sticky="nsew")
        self.navigation_frame.grid_rowconfigure(4, weight=1)

        self.navigation_frame_label = customtkinter.CTkLabel(self.navigation_frame, text="  uTranscribe", image=self.logo_image,
                                                             compound="left", font=customtkinter.CTkFont(size=15, weight="bold"))
        self.navigation_frame_label.grid(row=0, column=0, padx=20, pady=20)

        self.home_button = customtkinter.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Home",
                                                   fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                   image=self.home_image, anchor="w", command=self.home_button_event)
        self.home_button.grid(row=1, column=0, sticky="ew")

        self.speech_button = customtkinter.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Speech Recognition",
                                                      fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                      image=self.speech_image, anchor="w", command=self.speech_button_event)
        self.speech_button.grid(row=2, column=0, sticky="ew")

        self.ocr_button = customtkinter.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Character Recognition",
                                                      fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                                      image=self.ocr_image, anchor="w", command=self.ocr_button_event)
        self.ocr_button.grid(row=3, column=0, sticky="ew")

        self.appearance_mode_menu = customtkinter.CTkOptionMenu(self.navigation_frame, values=["Light", "Dark"],
                                                                command=self.change_appearance_mode_event)
        self.appearance_mode_menu.grid(row=6, column=0, padx=20, pady=20, sticky="s")

        # create home frame
        self.home_frame = customtkinter.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.home_frame.grid_columnconfigure(0, weight=1)

        self.home_label = customtkinter.CTkLabel(self.home_frame, text="Welcome to \nuTranscribe!",
                                                  font=customtkinter.CTkFont(size=30, weight="bold"))
        self.home_label.grid(row=1, column=0, padx=20, pady=10)
        tabview_1 = customtkinter.CTkTabview(self.home_frame, width=1280, height=300)
        tabview_1.grid(pady=10, padx=10)
        tabview_1.add("Speech Recognition")
        label = customtkinter.CTkLabel(master=tabview_1,
                               text="This feature will allow you to convert Speech from videos to a text format. \n\nUpload video button lets the user upload a video from the PC. \n\n Transcribe button lets you transcribe the file you chose. \n\nSave as function lets you save your file to a .txt, .docx or .pdf format.",
                               width=10,
                               height=10,
                               font=customtkinter.CTkFont(size=18, weight="normal"))
        label.place(relx= 0.5, rely=0.5, anchor=tkinter.CENTER)
        
        tabview_2 = customtkinter.CTkTabview(self.home_frame, width=1280, height=200)
        tabview_2.grid(pady=10, padx=10)
        tabview_2.add("Character Recognition")
        label = customtkinter.CTkLabel(master=tabview_2,
                               text="This feature will allow you to select a part of your \nscreen to take a screenshot and extract the text from it. \n\nCapture button lets the user take a screenshot. \n\nSave as function lets you save your file to a .txt, .docx or .pdf format.",
                               width=10,
                               height=10,
                               font=customtkinter.CTkFont(size=18, weight="normal"))
        label.place(relx= 0.5, rely=0.5, anchor=tkinter.CENTER)
        # create speech frame
        self.speech_frame = customtkinter.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.speechbutton_frame = customtkinter.CTkFrame(master=self.speech_frame)
        self.speechbutton_frame.grid(row=1,column=1,padx=10,pady=10)
        
        self.speech_frame.grid_columnconfigure(0, weight=1)
        self.speech_frame_button_upload=customtkinter.CTkButton(self.speechbutton_frame, text="Upload Video or audio", image=self.image_icon_image, compound="top",command=convert)
        self.speech_frame_button_upload.grid(row=1, column=0, padx=0, pady=10)
        self.speech_frame_button_trans = customtkinter.CTkButton(self.speechbutton_frame, text="Transcribe Video", command=transcribe)
        self.speech_frame_button_trans.grid(row=2, column=0, padx=20, pady=10)
        self.speech_frame_textbox = customtkinter.CTkTextbox(self.speech_frame, height=500)
        self.speech_frame_textbox.grid(row=1, column=0, padx=(20, 20), pady=(20, 20), sticky="nsew")
        self.speech_frame_save_as = customtkinter.CTkButton(self.speechbutton_frame, text="Save As .txt", command=save1txt)
        self.speech_frame_save_as.grid(row=3, column=0, padx=20, pady=10)
        self.speech_frame_save_as_docx = customtkinter.CTkButton(self.speechbutton_frame, text="Save As .docx", command=save1docx)
        self.speech_frame_save_as_docx.grid(row=4, column=0, padx=20, pady=10)
        self.speech_frame_save_as_pdf = customtkinter.CTkButton(self.speechbutton_frame, text="Save As .pdf", command=save1pdf)
        self.speech_frame_save_as_pdf.grid(row=5, column=0, padx=20, pady=10)
        self.filelabel = customtkinter.CTkLabel(self.speech_frame, text="Select a File:", font=customtkinter.CTkFont(size=20, weight="normal"))
        self.filelabel.grid(row=0, column=0, padx=10, pady=10)
        
        self.ocr_frame = customtkinter.CTkFrame(self, corner_radius=0, fg_color="transparent")
        self.ocr_frame.grid_columnconfigure(0, weight=1)
        self.ocrbutton_frame = customtkinter.CTkFrame(master=self.ocr_frame)
        self.ocrbutton_frame.grid(row=1,column=1,padx=10,pady=10)
        self.ocr_frame_button_upload=customtkinter.CTkButton(self.ocrbutton_frame, text="Capture Image", image=self.image_icon_image, compound="top",command=snip)
        self.ocr_frame_button_upload.grid(row=1, column=0, padx=20, pady=10)
        self.toplevel_window = None 
        self.ocr_frame_textbox = customtkinter.CTkTextbox(self.ocr_frame, height=500)
        self.ocr_frame_textbox.grid(row=1, column=0, padx=(20, 20), pady=(20, 20), sticky="nsew")
        self.ocr_frame_save_as = customtkinter.CTkButton(self.ocrbutton_frame, text="Save As .txt", command=save2)
        self.ocr_frame_save_as.grid(row=2, column=0, padx=20, pady=10)
        self.ocr_frame_save_as = customtkinter.CTkButton(self.ocrbutton_frame, text="Save As .docx", command=save2docx)
        self.ocr_frame_save_as.grid(row=3, column=0, padx=20, pady=10)
        self.ocr_frame_save_as = customtkinter.CTkButton(self.ocrbutton_frame, text="Save As .pdf", command=save2pdf)
        self.ocr_frame_save_as.grid(row=4, column=0, padx=20, pady=10)
        # select default frame
        self.select_frame_by_name("home")

    def select_frame_by_name(self, name):
        # set button color for selected button
        self.home_button.configure(fg_color=("gray75", "gray25") if name == "home" else "transparent")
        self.speech_button.configure(fg_color=("gray75", "gray25") if name == "speech" else "transparent")
        self.ocr_button.configure(fg_color=("gray75", "gray25") if name == "ocr" else "transparent")

        # show selected frame
        if name == "home":
            self.home_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.home_frame.grid_forget()
        if name == "speech":
            self.speech_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.speech_frame.grid_forget()
        if name == "ocr":
            self.ocr_frame.grid(row=0, column=1, sticky="nsew")
        else:
            self.ocr_frame.grid_forget()

    def home_button_event(self):
        self.select_frame_by_name("home")

    def speech_button_event(self):
        self.select_frame_by_name("speech")

    def ocr_button_event(self):
        self.select_frame_by_name("ocr")
    
    def savetextbox(self, transcription):
        self.speech_frame_textbox.insert("0.0",transcription)

    def savetextbox2(self, transcription):
        self.ocr_frame_textbox.insert("0.0",transcription)
         
    def change_appearance_mode_event(self, new_appearance_mode):
        customtkinter.set_appearance_mode(new_appearance_mode)
    
if __name__ == "__main__":
    app = App()
    app.mainloop()